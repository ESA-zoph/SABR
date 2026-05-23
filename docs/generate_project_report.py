from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
FIGURES_DIR = DOCS_DIR / "figures"
SEED_TABLE = ROOT / "data" / "training" / "repeats_cas_types_seed.csv"
CANDIDATE_TABLE = ROOT / "data" / "training" / "repeats_cas_types_augmented_vink_genbank_targeted.csv"
REPORT_PATH = DOCS_DIR / "SABR_model_development_report.docx"
MODEL_COMPARISON = DOCS_DIR / "model_comparison_current.csv"
BEST_MODEL_PREDICTIONS = DOCS_DIR / "best_model_predictions.csv"


def main() -> None:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    seed = _read_table(SEED_TABLE)
    candidate = _read_table(CANDIDATE_TABLE)
    model_comparison = _read_table(MODEL_COMPARISON)

    type_figure = FIGURES_DIR / "cas_type_distribution.png"
    subtype_figure = FIGURES_DIR / "cas_subtype_distribution.png"
    model_figure = FIGURES_DIR / "model_validation_plan.png"
    accuracy_figure = FIGURES_DIR / "model_accuracy_comparison.png"
    confusion_figure = FIGURES_DIR / "best_model_confusion_matrix.png"
    f1_figure = FIGURES_DIR / "best_model_per_class_f1.png"
    roc_figure = FIGURES_DIR / "best_model_roc_curve.png"
    importance_figure = FIGURES_DIR / "best_model_feature_importance.png"
    error_subtype_figure = FIGURES_DIR / "best_model_error_by_subtype.png"
    confidence_figure = FIGURES_DIR / "best_model_confidence_correct_vs_wrong.png"
    top_errors_figure = FIGURES_DIR / "best_model_top_errors.png"
    _plot_type_distribution(candidate, type_figure)
    _plot_subtype_distribution(candidate, subtype_figure)
    _plot_model_validation_plan(model_figure)

    document = Document()
    document.add_heading("SABR Model Development Report", level=0)
    document.add_paragraph(
        "SABR (Spacer Alignment-Based Resistance) is being developed as a "
        "transparent CRISPR-phage targeting evidence tool. This report tracks "
        "the repeat-to-Cas-type model work and the dataset-building strategy."
    )

    document.add_heading("Scientific Goal", level=1)
    document.add_paragraph(
        "The model goal is to predict likely CRISPR-Cas type or subtype from "
        "CRISPR repeat and array features when complete Cas annotation is not "
        "available. The output should be probabilistic evidence, not a definitive "
        "biological claim."
    )

    document.add_heading("Dataset Layers", level=1)
    document.add_paragraph(
        "Two dataset layers are being kept separate because they have different "
        "levels of label confidence."
    )
    _add_dataset_table(document, seed, "Manual curated seed set", "curated or literature-supported labels")
    _add_dataset_table(
        document,
        candidate,
        "Scaled computational candidate set",
        "CRISPRCasdb/Vink 2021 rows filtered by repeat/proximity subtype agreement",
    )

    document.add_heading("Current Candidate Dataset Composition", level=1)
    document.add_paragraph(
        f"The current scaled candidate table contains {len(candidate):,} rows. "
        "It is useful for model development, but it should not be presented as "
        "a manually curated gold standard."
    )
    document.add_picture(str(type_figure), width=Inches(5.8))
    document.add_paragraph("Figure 1. Candidate rows by broad CRISPR-Cas type.")
    document.add_picture(str(subtype_figure), width=Inches(6.2))
    document.add_paragraph("Figure 2. Candidate rows by CRISPR-Cas subtype.")

    document.add_heading("Feature Set", level=1)
    for item in [
        "repeat sequence",
        "repeat length",
        "repeat GC percent",
        "repeat k-mer frequencies",
        "spacer count",
        "mean spacer length",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_paragraph(
        "Organism, taxonomy, and accession metadata are retained for auditing "
        "and split design, but they should not be used as first-model features "
        "because they can create shortcut learning."
    )

    document.add_heading("Model Ladder", level=1)
    document.add_picture(str(model_figure), width=Inches(6.2))
    document.add_paragraph("Figure 3. Planned model and validation ladder.")
    for item in [
        "Nearest-repeat similarity baseline for interpretability.",
        "Regularized logistic regression or linear SVM for a simple supervised baseline.",
        "Random forest or gradient boosting for non-linear repeat/array feature interactions.",
        "Sequence CNN or transformer-style model only after enough high-confidence rows exist.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Current Smoke-Test Results", level=1)
    document.add_paragraph(
        "The current model comparison evaluates simple and moderately complex "
        "methods on the augmented Vink/CRISPRCasdb plus targeted GenBank table. "
        "Random row-level splits remain smoke tests; genome/accession and genus "
        "holdouts are the more relevant validation steps."
    )
    if not model_comparison.empty:
        _add_model_comparison_table(document, model_comparison)
    for figure_path, caption in [
        (accuracy_figure, "Figure 4. Accuracy comparison across validation strategies and models."),
        (confusion_figure, "Figure 5. Confusion matrix for the current best ExtraTrees predictor."),
        (f1_figure, "Figure 6. Per-class F1 scores for the current best predictor."),
        (roc_figure, "Figure 7. One-vs-rest ROC curves for probability-capable best predictor."),
        (importance_figure, "Figure 8. Top feature importances for the current best predictor."),
        (error_subtype_figure, "Figure 9. Best-model error rate by subtype in the held-out genus split."),
        (confidence_figure, "Figure 10. Best-model confidence distribution for correct versus wrong calls."),
        (top_errors_figure, "Figure 11. Most frequent subtype-level confusions from the best predictor."),
    ]:
        if figure_path.exists():
            document.add_picture(str(figure_path), width=Inches(6.2))
            document.add_paragraph(caption)
    if BEST_MODEL_PREDICTIONS.exists():
        predictions = _read_table(BEST_MODEL_PREDICTIONS)
        if not predictions.empty and "correct" in predictions.columns:
            correct = int(predictions["correct"].sum())
            total = int(len(predictions))
            document.add_paragraph(
                "The held-out prediction export is stored at "
                f"`docs/best_model_predictions.csv` and currently contains "
                f"{total:,} genus-holdout predictions, with {correct:,} correct "
                f"and {total - correct:,} wrong calls."
            )

    document.add_heading("Validation Plan", level=1)
    for item in [
        "Random row split: use only as a quick smoke test.",
        "Genome/accession holdout: required next validation step.",
        "Species holdout: tests generalization beyond strains and close relatives.",
        "Genus holdout: strongest practical test for broad deployment claims.",
        "Source holdout: test whether performance survives when an entire source is excluded from training.",
        "Calibration analysis: report whether prediction confidence is reliable.",
        "Error analysis: identify which subtypes are confused and whether rare classes need merging or more data.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Immediate Next Work", level=1)
    for item in [
        "Implement accession/species/genus holdout evaluation in the training command.",
        "Add more targeted Type V and Type VI rows because they remain underrepresented.",
        "Keep manually curated gold rows separate from computational candidate rows.",
        "Add model comparison tables for nearest repeat, logistic regression/SVM, random forest, and gradient boosting.",
        "Document all data sources and label-confidence tiers before making publication claims.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("References and Sources", level=1)
    for item in [
        "Vink et al. 2021, Genome Biology, CRISPRCasdb-derived repeat/PAM supplementary data.",
        "PA14 Type I-F CRISPR-phage literature sources tracked in data/curation/literature_sources.tsv.",
        "E. coli K-12 Type I-E review source tracked in data/curation/literature_sources.tsv.",
        "S. pyogenes SF370 Type II-A, F. novicida U112 Type V-A, and S. epidermidis RP62A Type III-A seed sources tracked in data/curation/literature_sources.tsv.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.save(REPORT_PATH)
    print(f"Wrote {REPORT_PATH}")


def _read_table(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def _add_dataset_table(document: Document, table: pd.DataFrame, name: str, confidence: str) -> None:
    document.add_heading(name, level=2)
    rows = len(table)
    genomes = table["genome_id"].nunique() if not table.empty else 0
    subtypes = table["cas_subtype"].nunique() if not table.empty else 0
    doc_table = document.add_table(rows=1, cols=4)
    headers = ["Rows", "Genomes/accessions", "Subtypes", "Label confidence"]
    for index, header in enumerate(headers):
        doc_table.rows[0].cells[index].text = header
    values = [f"{rows:,}", f"{genomes:,}", f"{subtypes:,}", confidence]
    row_cells = doc_table.add_row().cells
    for index, value in enumerate(values):
        row_cells[index].text = value


def _add_model_comparison_table(document: Document, table: pd.DataFrame) -> None:
    visible = table[["split_strategy", "method", "train_rows", "test_rows", "accuracy"]].copy()
    doc_table = document.add_table(rows=1, cols=len(visible.columns))
    for index, column in enumerate(visible.columns):
        doc_table.rows[0].cells[index].text = column
    for _, row in visible.iterrows():
        cells = doc_table.add_row().cells
        for index, column in enumerate(visible.columns):
            value = row[column]
            cells[index].text = f"{value:.4f}" if column == "accuracy" else str(value)


def _plot_type_distribution(table: pd.DataFrame, output_path: Path) -> None:
    counts = table["cas_type"].value_counts().sort_index()
    _bar_plot(counts, output_path, "Candidate Rows by CRISPR-Cas Type", "Type", "Rows")


def _plot_subtype_distribution(table: pd.DataFrame, output_path: Path) -> None:
    counts = table["cas_subtype"].value_counts().sort_values(ascending=True)
    fig_height = max(4.0, 0.28 * len(counts))
    plt.figure(figsize=(8.5, fig_height))
    plt.barh(counts.index, counts.values, color="#2F6F73")
    plt.xlabel("Rows")
    plt.ylabel("Subtype")
    plt.title("Candidate Rows by CRISPR-Cas Subtype")
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()


def _plot_model_validation_plan(output_path: Path) -> None:
    labels = [
        "Nearest\nrepeat",
        "Linear\nmodel",
        "Random forest /\nboosting",
        "Sequence\nmodel",
    ]
    validation = [
        "Row split",
        "Accession\nholdout",
        "Species\nholdout",
        "Genus\nholdout",
    ]
    plt.figure(figsize=(8, 3.5))
    for index, label in enumerate(labels):
        plt.scatter(index, 1, s=1800, color="#1F4E5F")
        plt.text(index, 1, label, ha="center", va="center", color="white", fontsize=9)
        if index < len(labels) - 1:
            plt.arrow(index + 0.22, 1, 0.56, 0, width=0.015, head_width=0.08, head_length=0.08, color="#5C6770")
    for index, label in enumerate(validation):
        plt.text(index, 0.18, label, ha="center", va="center", fontsize=9)
    plt.text(1.5, 1.55, "Model complexity increases only if validation supports it", ha="center", fontsize=11)
    plt.text(1.5, -0.05, "Validation stringency increases from smoke test to publication-grade", ha="center", fontsize=10)
    plt.xlim(-0.6, 3.6)
    plt.ylim(-0.25, 1.85)
    plt.axis("off")
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()


def _bar_plot(counts: pd.Series, output_path: Path, title: str, xlabel: str, ylabel: str) -> None:
    plt.figure(figsize=(7, 4))
    plt.bar(counts.index, counts.values, color="#2F6F73")
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()


if __name__ == "__main__":
    main()

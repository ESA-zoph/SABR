from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "manuscript_assets"
SOURCE = DOCS / "SABR_manuscript_draft.md"
OUTPUT = DOCS / "SABR_manuscript.docx"

FIGURES = [
    ("targeting_score_heatmap.png", "Figure 1. Bacteria-by-phage SABR targeting-evidence heatmap; values are evidence scores, not resistance phenotypes."),
    ("benchmark_score_summary.png", "Figure 2. Benchmark CRISPR targeting evidence scores and the PAM/PFS-unsupported score cap."),
    ("architecture_selection_roc.png", "Figure 3. One-vs-rest ROC comparison of probability-capable subtype architectures evaluated on the same genus-held-out development split."),
    ("model_comparison_summary.png", "Figure 4. ExtraTrees development analyses across annotation-table and grouped-validation designs; split designs differ."),
    ("typeiii_performance_summary.png", "Figure 5. Type III subtype F1 scores during SABR model development."),
    ("selected_model_roc_focus_subtypes.png", "Figure 6. Genome-held-out ROC curves for the selected SABR model, highlighting difficult Type III and adjacent Type I classes."),
    ("selected_model_confusion_heatmap.png", "Figure 7. Row-normalized held-out confusion heatmap for the selected SABR subtype model."),
    ("selected_model_tsne_group_and_subtype_panels.png", "Figure 8. Held-out t-SNE projection by broad Cas type and by the Type III/adjacent Type I error region; projection is descriptive only."),
    ("calibration_comparison.png", "Figure 9. Confidence calibration of the earlier development-stage and selected SABR subtype models."),
    ("selected_model_builtin_feature_importance.png", "Figure 10. Highest-ranking FASTA-derived features in the selected SABR ExtraTrees subtype model."),
    ("selected_model_feature_category_importance.png", "Figure 11. Feature importance aggregated by biological feature family."),
    ("selected_model_top_error_pairs.png", "Figure 12. Most frequent held-out subtype confusions for the selected SABR model."),
    ("cctyper_pilot_summary.png", "Figure 13. First independent CCTyper-supported pilot evaluation; enriched strict-validation results are reported in the text."),
]


def main() -> None:
    document = Document()
    _render_markdown(document, SOURCE.read_text(encoding="utf-8"))
    document.add_heading("Figures", level=1)
    for filename, caption in FIGURES:
        _add_figure(document, ASSETS / filename, caption)
    document.add_paragraph(
        "Matching SVG files are stored in docs/manuscript_assets/ for publication-quality editing."
    )
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


def _render_markdown(document: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and _is_separator(lines[index + 1]):
            table_lines = [line, lines[index + 1].strip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            if level == 1:
                document.add_heading(title, 0)
            else:
                document.add_heading(title, min(level - 1, 3))
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line[:3].strip(".").isdigit() and ". " in line[:4]:
            document.add_paragraph(line.split(". ", 1)[1], style="List Number")
        elif line.startswith("Equation "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line.replace("`", ""))
            run.italic = True
        else:
            paragraph = line
            while index + 1 < len(lines):
                next_line = lines[index + 1].strip()
                if not next_line or next_line.startswith(("#", "-", "|")):
                    break
                if next_line[:3].strip(".").isdigit() and ". " in next_line[:4]:
                    break
                paragraph += " " + next_line
                index += 1
            document.add_paragraph(paragraph.replace("**", "").replace("`", ""))
        index += 1


def _is_separator(line: str) -> bool:
    stripped = line.strip().replace("|", "").replace(":", "").replace("-", "").strip()
    return not stripped


def _add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in lines if not _is_separator(line)]
    table = document.add_table(rows=1, cols=len(rows[0]))
    for column, value in enumerate(rows[0]):
        table.rows[0].cells[column].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for column, value in enumerate(row):
            cells[column].text = value


def _add_figure(document: Document, path: Path, caption: str) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(6.2))
        document.add_paragraph(caption)


if __name__ == "__main__":
    main()

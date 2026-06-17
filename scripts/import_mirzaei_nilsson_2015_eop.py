from __future__ import annotations

import argparse
from pathlib import Path
import re
import sys

import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from crispr_phage_predictor.interactions import (
    INTERACTION_COLUMNS,
    eop_class_from_value,
    parse_eop,
    susceptibility_from_eop_class,
    validate_interaction_table,
)


PHAGES = ["SU10", "SU16", "SU27", "SU32", "SU57", "SU63"]
SECTION_TO_BACTERIUM = {
    "ECOR": "Escherichia coli",
    "ESBL E. coli": "Escherichia coli",
    "SARA": "Salmonella enterica",
    "SARB": "Salmonella enterica",
}


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Import Mirzaei and Nilsson 2015 EOP supplementary table."
    )
    parser.add_argument("docx_path", type=Path)
    parser.add_argument("output_tsv", type=Path)
    args = parser.parse_args()

    rows = import_rows(args.docx_path)
    table = pd.DataFrame(rows, columns=INTERACTION_COLUMNS)
    validate_interaction_table(table)
    args.output_tsv.parent.mkdir(parents=True, exist_ok=True)
    table.to_csv(args.output_tsv, sep="\t", index=False)


def import_rows(docx_path: Path) -> list[dict[str, str]]:
    import docx

    document = docx.Document(docx_path)
    if len(document.tables) != 1:
        raise ValueError(f"Expected one DOCX table, found {len(document.tables)}")

    rows: list[dict[str, str]] = []
    section = ""
    for table_row in document.tables[0].rows[1:]:
        cells = [_cell_text(cell.text) for cell in table_row.cells]
        first = cells[0]
        if first in SECTION_TO_BACTERIUM:
            section = first
            continue
        if not section or not first:
            continue
        strain = _strain_name(section, first)
        bacterium = SECTION_TO_BACTERIUM[section]
        for phage_index, phage in enumerate(PHAGES, start=1):
            raw_eop = cells[phage_index]
            if not raw_eop:
                continue
            relation, value = parse_eop(_strip_footnote(raw_eop))
            eop_class = eop_class_from_value(value, relation)
            rows.append(
                {
                    "interaction_id": _interaction_id(section, first, phage),
                    "source_key": "MirzaeiNilsson2015_EOP",
                    "source_type": "supplement",
                    "pmid": "25794006",
                    "doi": "10.1371/journal.pone.0118557",
                    "assay_type": "eop",
                    "bacterium": bacterium,
                    "strain": strain,
                    "bacterial_accession": "",
                    "phage": phage,
                    "phage_accession": "",
                    "reference_host": "primary isolation host",
                    "raw_eop": _strip_footnote(raw_eop),
                    "eop_relation": relation,
                    "eop_value": "" if value is None else f"{value:.8g}",
                    "eop_class": eop_class,
                    "susceptibility_label": susceptibility_from_eop_class(eop_class),
                    "plaque_result": _plaque_result(eop_class),
                    "anti_crispr_status": "not_evaluated",
                    "anti_crispr_genes": "",
                    "crispr_interference_evidence": "not_evaluated",
                    "other_defense_evidence": _other_defense_note(cells),
                    "experimental_conditions": (
                        "Stationary-phase bacteria, double-layer plaque assay, "
                        "30 C overnight; EOP is average PFU on target divided by "
                        "average PFU on host, three measurements."
                    ),
                    "curation_status": "curated",
                    "curation_confidence": "high",
                    "notes": (
                        "Imported from S1 Table; blank source cells were not treated "
                        "as measured interactions."
                    ),
                }
            )
    return rows


def _cell_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def _strip_footnote(value: str) -> str:
    value = _cell_text(value)
    return re.sub(r"(?<=\d)[a-d]$", "", value)


def _strain_name(section: str, strain_id: str) -> str:
    prefix = {
        "ECOR": "ECOR",
        "ESBL E. coli": "ESBL_Ecoli",
        "SARA": "SARA",
        "SARB": "SARB",
    }[section]
    return f"{prefix}_{strain_id}"


def _interaction_id(section: str, strain_id: str, phage: str) -> str:
    safe_section = re.sub(r"[^a-z0-9]+", "_", section.lower()).strip("_")
    safe_strain = re.sub(r"[^a-z0-9]+", "_", strain_id.lower()).strip("_")
    return f"mirzaei2015_{safe_section}_{safe_strain}_{phage.lower()}_eop"


def _plaque_result(eop_class: str) -> str:
    if eop_class == "none":
        return "no_plaques"
    if eop_class in {"trace", "low"}:
        return "pinpoint_plaques"
    if eop_class in {"medium", "high"}:
        return "clear_plaques"
    return "not_reported"


def _other_defense_note(cells: list[str]) -> str:
    notes = []
    if len(cells) > 7 and cells[7]:
        notes.append(f"colicin/plasmid marker: {cells[7]}")
    if len(cells) > 8 and cells[8]:
        notes.append(f"P2 prophage marker: {cells[8]}")
    return "; ".join(notes)


if __name__ == "__main__":
    main()
